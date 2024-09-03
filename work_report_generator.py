import os
from docx import Document
from docx.shared import Pt
from datetime import datetime
import tkinter as tk
from tkinter import messagebox, scrolledtext

def create_work_report(full_name, date, work_plan, status_of_completion):
    # Create a new Document
    doc = Document()
    
    # Title
    doc.add_heading('Daily Work Report', 0)
    
    # Create a table with 2 columns and 4 rows
    table = doc.add_table(rows=4, cols=2)
    
    # Define column widths (if desired; adjust as needed)
    for row in table.rows:
        row.cells[0].width = Pt(100)
        row.cells[1].width = Pt(300)
    
    # Populate the table
    table.cell(0, 0).text = 'Full Name:'
    table.cell(0, 1).text = full_name
    
    table.cell(1, 0).text = 'Date:'
    table.cell(1, 1).text = date
    
    table.cell(2, 0).text = 'Work Plan:'
    table.cell(2, 1).text = work_plan
    
    table.cell(3, 0).text = 'Status of Completion:'
    table.cell(3, 1).text = status_of_completion
    
    # Save the document to Desktop
    desktop_path = os.path.expanduser("~/Desktop")
    file_name = f"{full_name.replace(' ', '_')}_Daily_Work_Report_{date.replace('/', '-')}.docx"
    full_path = os.path.join(desktop_path, file_name)
    doc.save(full_path)
    
    messagebox.showinfo("Success", f"Report saved as {full_path}")

    # Show the Exit button after the report is generated
    exit_btn.pack(pady=10)

# Function to get data from the GUI and generate the report
def generate_report():
    full_name = name_entry.get()
    work_plan = work_plan_text.get("1.0", tk.END).strip()
    status_of_completion = status_text.get("1.0", tk.END).strip()

    if not full_name or not work_plan or not status_of_completion:
        messagebox.showerror("Error", "Please fill out all fields")
        return
    
    current_date = datetime.now().strftime("%d/%m/%Y")
    create_work_report(full_name, current_date, work_plan, status_of_completion)

# Function to exit the application
def exit_app():
    root.quit()

# Create the GUI window
root = tk.Tk()
root.title("Daily Work Report Generator")
root.geometry("400x450")

# Labels and Entry Fields
tk.Label(root, text="Full Name").pack(pady=5)
name_entry = tk.Entry(root, width=40)
name_entry.pack(pady=5)

tk.Label(root, text="Work Plan").pack(pady=5)
work_plan_text = scrolledtext.ScrolledText(root, wrap=tk.WORD, width=40, height=5)
work_plan_text.pack(pady=5)

tk.Label(root, text="Status of Completion").pack(pady=5)
status_text = scrolledtext.ScrolledText(root, wrap=tk.WORD, width=40, height=5)
status_text.pack(pady=5)

# Generate Report Button
generate_btn = tk.Button(root, text="Generate Report", command=generate_report)
generate_btn.pack(pady=20)

# Exit Button (Initially hidden)
exit_btn = tk.Button(root, text="Exit", command=exit_app)
# exit_btn.pack() is called only after generating the report

# Start the GUI loop
root.mainloop()
